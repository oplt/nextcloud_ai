from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from email import policy
from email.parser import BytesParser
from email.utils import getaddresses, parsedate_to_datetime
import html
import csv
from pathlib import Path
import xml.etree.ElementTree as ET
from zipfile import BadZipFile, ZipFile

import docx
import pdfplumber


@dataclass(slots=True)
class ParsedPage:
    page_number: int | None
    text: str


@dataclass(slots=True)
class ParsedDocument:
    text: str
    pages: list[ParsedPage]
    metadata: dict[str, object]
    attachments: list["ParsedAttachment"] = field(default_factory=list)


@dataclass(slots=True)
class ParsedAttachment:
    file_name: str
    mime_type: str | None
    payload: bytes
    is_inline: bool = False
    content_id: str | None = None


class UnsupportedDocumentTypeError(ValueError):
    pass


PDF_MIME_TYPES = {"application/pdf"}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
}
ODT_MIME_TYPES = {"application/vnd.oasis.opendocument.text"}
TEXT_MIME_TYPES = {
    "text/plain",
    "text/markdown",
    "text/x-markdown",
    "application/octet-stream",
}
CSV_MIME_TYPES = {"text/csv", "application/csv"}
IMAGE_MIME_TYPES = {
    "image/jpeg",
    "image/png",
    "image/gif",
    "image/webp",
    "image/bmp",
    "image/tiff",
    "image/svg+xml",
}
EMAIL_MIME_TYPES = {"message/rfc822", "application/eml"}

_AMOUNT_RE = re.compile(
    r"(?i)(?:total|amount|balance|subtotal|btw|vat|totaal|bedrag|factuurbedrag)"
    r"[^0-9€$]{0,30}([€$]?\s?\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?)"
)

_DATE_RE = re.compile(
    r"\b(?:\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\d{4}[-/]\d{1,2}[-/]\d{1,2})\b"
)

_INVOICE_NO_RE = re.compile(
    r"(?i)(?:invoice|factuur|invoice no|factuurnummer|nummer|nr\.?)"
    r"[^A-Z0-9]{0,20}([A-Z0-9][A-Z0-9\-_/]{2,})"
)

ODT_OFFICE_NS = "urn:oasis:names:tc:opendocument:xmlns:office:1.0"
ODT_TEXT_NS = "urn:oasis:names:tc:opendocument:xmlns:text:1.0"
ODT_DOCUMENT_CONTENT_TAG = f"{{{ODT_OFFICE_NS}}}document-content"
ODT_PARAGRAPH_TAG = f"{{{ODT_TEXT_NS}}}p"
ODT_HEADING_TAG = f"{{{ODT_TEXT_NS}}}h"
ODT_LINE_BREAK_TAG = f"{{{ODT_TEXT_NS}}}line-break"
ODT_TAB_TAG = f"{{{ODT_TEXT_NS}}}tab"
ODT_SPACE_TAG = f"{{{ODT_TEXT_NS}}}s"


async def parse_document_bytes(
    file_name: str, mime_type: str | None, payload: bytes
) -> ParsedDocument:
    suffix = Path(file_name).suffix.lower()
    normalized_mime = (mime_type or "").lower()

    if suffix == ".pdf" or normalized_mime in PDF_MIME_TYPES:
        return parse_pdf_bytes(payload)
    if suffix == ".docx" or normalized_mime in DOCX_MIME_TYPES:
        return parse_docx_bytes(payload)
    if suffix == ".odt" or normalized_mime in ODT_MIME_TYPES:
        return parse_odt_bytes(payload)
    if suffix == ".eml" or normalized_mime in EMAIL_MIME_TYPES:
        return parse_email_bytes(payload)
    if suffix in {
        ".jpg",
        ".jpeg",
        ".png",
        ".gif",
        ".webp",
        ".bmp",
        ".tif",
        ".tiff",
        ".svg",
    } or normalized_mime in IMAGE_MIME_TYPES:
        return parse_image_bytes(file_name=file_name, mime_type=mime_type, payload=payload)
    if suffix == ".csv" or normalized_mime in CSV_MIME_TYPES:
        return parse_csv_bytes(payload)
    if suffix in {".txt", ".md", ".markdown"} or normalized_mime in TEXT_MIME_TYPES:
        return parse_text_bytes(
            payload,
            markdown=suffix in {".md", ".markdown"} or "markdown" in normalized_mime,
        )
    raise UnsupportedDocumentTypeError(f"Unsupported document type for {file_name}")


def parse_pdf_bytes(payload: bytes) -> ParsedDocument:
    pages: list[ParsedPage] = []
    with pdfplumber.open(io.BytesIO(payload)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            page_parts: list[str] = []
            text = (page.extract_text() or "").strip()
            if text:
                page_parts.append(text)
            for table in page.extract_tables() or []:
                table_text = _format_table(table)
                if table_text:
                    page_parts.append(table_text)
            if page_parts:
                pages.append(ParsedPage(page_number=index, text="\n\n".join(page_parts)))
    combined = "\n\n".join(page.text for page in pages)
    return ParsedDocument(
        text=combined,
        pages=pages,
        metadata={"page_count": len(pages),
                  "parser": "pdfplumber",
                  "extracted_fields": extract_generic_financial_fields(text),},
    )


def parse_docx_bytes(payload: bytes) -> ParsedDocument:
    document = docx.Document(io.BytesIO(payload))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if not paragraph_text:
            continue
        style_name = (paragraph.style.name if paragraph.style is not None else "").lower()
        if style_name.startswith("heading"):
            level = _heading_level_from_style(style_name)
            blocks.append(f"{'#' * level} {paragraph_text}")
        else:
            blocks.append(paragraph_text)
    for table in document.tables:
        table_text = _format_table(
            [[cell.text.strip() for cell in row.cells] for row in table.rows]
        )
        if table_text:
            blocks.append(table_text)
    text = "\n\n".join(blocks)
    pages = [ParsedPage(page_number=None, text=text)] if text else []
    return ParsedDocument(
        text=text,
        pages=pages,
        metadata={
            "block_count": len(blocks),
            "table_count": len(document.tables),
            "parser": "python-docx",
            "extracted_fields": extract_generic_financial_fields(text),
        },
    )


def parse_odt_bytes(payload: bytes) -> ParsedDocument:
    try:
        with ZipFile(io.BytesIO(payload)) as archive:
            content_xml = archive.read("content.xml")
    except (BadZipFile, KeyError) as exc:
        raise ValueError("Invalid ODT payload") from exc

    try:
        root = ET.fromstring(content_xml)
    except ET.ParseError as exc:
        raise ValueError("Invalid ODT XML payload") from exc

    if root.tag != ODT_DOCUMENT_CONTENT_TAG:
        raise ValueError("Invalid ODT document root")

    document_text = root.find(f"./{{{ODT_OFFICE_NS}}}body/{{{ODT_OFFICE_NS}}}text")
    if document_text is None:
        raise ValueError("Invalid ODT document body")

    blocks = [
        block_text
        for element in document_text.iter()
        if element.tag in {ODT_PARAGRAPH_TAG, ODT_HEADING_TAG}
        if (block_text := _extract_odt_text(element))
    ]
    text = "\n".join(blocks)
    pages = [ParsedPage(page_number=None, text=text)] if text else []
    return ParsedDocument(
        text=text,
        pages=pages,
        metadata={"block_count": len(blocks), "parser": "odt-xml", "extracted_fields": extract_generic_financial_fields(text),},

    )


def parse_text_bytes(payload: bytes, *, markdown: bool = False) -> ParsedDocument:
    text = _decode_text(payload)
    pages = [ParsedPage(page_number=None, text=text)] if text.strip() else []
    return ParsedDocument(
        text=text,
        pages=pages,
        metadata={
            "parser": "plain-text",
            "markdown": markdown,
            "line_count": len(text.splitlines()),
            "extracted_fields": extract_generic_financial_fields(text),
        },
    )


def parse_csv_bytes(payload: bytes) -> ParsedDocument:
    text = _decode_text(payload)
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return ParsedDocument(text="", pages=[], metadata={"parser": "csv", "row_count": 0})
    formatted = _format_table(rows[:5000])
    pages = [ParsedPage(page_number=None, text=formatted)] if formatted.strip() else []
    return ParsedDocument(
        text=formatted,
        pages=pages,
        metadata={"parser": "csv", "row_count": len(rows), "truncated": len(rows) > 5000, "extracted_fields": extract_generic_financial_fields(text),},
    )


def parse_email_bytes(payload: bytes) -> ParsedDocument:
    message = BytesParser(policy=policy.default).parsebytes(payload)
    body = _extract_email_body(message)
    attachments = _extract_email_attachments(message)

    subject = _clean_header(message.get("Subject"))
    sender = _format_addresses(message.get_all("From", []))
    recipients = _format_addresses(
        [*message.get_all("To", []), *message.get_all("Cc", [])]
    )
    sent_at = _normalize_email_datetime(message.get("Date"))
    message_id = _clean_header(message.get("Message-ID"))
    references = [
        ref
        for ref in re.split(r"\s+", _clean_header(message.get("References")))
        if ref
    ]
    in_reply_to = _clean_header(message.get("In-Reply-To"))
    thread_key = references[0] if references else (in_reply_to or message_id)

    attachment_lines = [
        f"- {attachment.file_name} ({attachment.mime_type or 'application/octet-stream'})"
        for attachment in attachments
    ]
    sections = [
        line
        for line in [
            f"Subject: {subject}" if subject else "",
            f"From: {sender}" if sender else "",
            f"To: {recipients}" if recipients else "",
            f"Date: {sent_at}" if sent_at else "",
            body.strip(),
            "Attachments:\n" + "\n".join(attachment_lines) if attachment_lines else "",
        ]
        if line
    ]
    text = "\n\n".join(sections).strip()
    pages = [ParsedPage(page_number=None, text=text)] if text else []
    return ParsedDocument(
        text=text,
        pages=pages,
        metadata={
            "parser": "email-rfc822",
            "subject": subject,
            "from": sender,
            "to": recipients,
            "date": sent_at,
            "message_id": message_id,
            "in_reply_to": in_reply_to,
            "references": references,
            "thread_key": thread_key,
            "attachment_count": len(attachments),
            "attachments": [
                {
                    "file_name": attachment.file_name,
                    "mime_type": attachment.mime_type,
                    "size_bytes": len(attachment.payload),
                    "is_inline": attachment.is_inline,
                }
                for attachment in attachments
            ],
            "extracted_fields": extract_generic_financial_fields(text),
        },
        attachments=attachments,
    )


def parse_image_bytes(
    *, file_name: str, mime_type: str | None, payload: bytes
) -> ParsedDocument:
    normalized_mime = (mime_type or "application/octet-stream").lower()
    text = (
        f"Image file: {Path(file_name).name}\n"
        f"MIME type: {normalized_mime}\n"
        "No OCR extracted text available for this image."
    )
    return ParsedDocument(
        text=text,
        pages=[ParsedPage(page_number=None, text=text)],
        metadata={
            "parser": "image-metadata-fallback",
            "ocr_applied": False,
            "size_bytes": len(payload),
            "mime_type": normalized_mime,
            "extracted_fields": extract_generic_financial_fields(text),
        },
    )


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="ignore")


def _extract_odt_text(element: ET.Element) -> str:
    parts: list[str] = []
    if element.text:
        parts.append(element.text)

    for child in element:
        if child.tag == ODT_LINE_BREAK_TAG:
            parts.append("\n")
        elif child.tag == ODT_TAB_TAG:
            parts.append("\t")
        elif child.tag == ODT_SPACE_TAG:
            count = int(child.attrib.get(f"{{{ODT_TEXT_NS}}}c", "1"))
            parts.append(" " * max(1, count))
        else:
            parts.append(_extract_odt_text(child))

        if child.tail:
            parts.append(child.tail)

    return "".join(parts).strip()


def _heading_level_from_style(style_name: str) -> int:
    match = re.search(r"(\d+)", style_name)
    if not match:
        return 2
    return max(1, min(6, int(match.group(1))))


def _format_table(rows) -> str:
    cleaned_rows: list[list[str]] = []
    for row in rows or []:
        cells = [str(cell or "").strip().replace("\n", " ") for cell in row]
        if any(cells):
            cleaned_rows.append(cells)
    if not cleaned_rows:
        return ""
    width = max(len(row) for row in cleaned_rows)
    normalized = [row + [""] * (width - len(row)) for row in cleaned_rows]
    table_lines = ["| " + " | ".join(row) + " |" for row in normalized]
    semantic_lines = _table_semantic_lines(normalized)
    if semantic_lines:
        return "\n".join([*table_lines, "", "Extracted table facts:", *semantic_lines])
    return "\n".join(table_lines)


def _table_semantic_lines(rows: list[list[str]]) -> list[str]:
    facts: list[str] = []
    if not rows:
        return facts

    for row_index, row in enumerate(rows, start=1):
        non_empty = [cell.strip() for cell in row if cell and cell.strip()]

        if len(non_empty) == 2:
            facts.append(f"- {non_empty[0]}: {non_empty[1]}")
            facts.append(f"- Field {non_empty[0]} has value {non_empty[1]}")

        if len(non_empty) > 2:
            facts.append(f"- Row {row_index}: " + "; ".join(non_empty))

    header = rows[0]
    if len(rows) >= 2 and any(header):
        for row_index, row in enumerate(rows[1:], start=1):
            pairs = [
                f"{heading.strip()}: {value.strip()}"
                for heading, value in zip(header, row)
                if heading and value and heading.strip() and value.strip()
            ]
            if pairs:
                facts.append(f"- Row {row_index}: " + "; ".join(pairs))

    return list(dict.fromkeys(facts))


def _extract_email_body(message) -> str:
    plain_parts: list[str] = []
    html_parts: list[str] = []

    if message.is_multipart():
        for part in message.walk():
            if part.is_multipart():
                continue
            disposition = (part.get_content_disposition() or "").lower()
            if disposition == "attachment":
                continue
            content_type = part.get_content_type().lower()
            text = _decode_message_part(part)
            if not text.strip():
                continue
            if content_type == "text/plain":
                plain_parts.append(text.strip())
            elif content_type == "text/html":
                html_parts.append(_strip_html(text).strip())
    else:
        text = _decode_message_part(message)
        if message.get_content_type().lower() == "text/html":
            html_parts.append(_strip_html(text).strip())
        else:
            plain_parts.append(text.strip())

    candidate_parts = plain_parts or html_parts
    return "\n\n".join(part for part in candidate_parts if part).strip()


def _extract_email_attachments(message) -> list[ParsedAttachment]:
    attachments: list[ParsedAttachment] = []
    for part in message.walk():
        if part.is_multipart():
            continue
        disposition = (part.get_content_disposition() or "").lower()
        file_name = part.get_filename()
        if disposition != "attachment" and not file_name:
            continue
        payload = part.get_payload(decode=True) or b""
        if not payload:
            continue
        attachments.append(
            ParsedAttachment(
                file_name=file_name or f"attachment-{len(attachments) + 1}",
                mime_type=part.get_content_type(),
                payload=payload,
                is_inline=disposition == "inline",
                content_id=_clean_header(part.get("Content-ID")),
            )
        )
    return attachments


def _decode_message_part(part) -> str:
    payload = part.get_payload(decode=True)
    if payload is None:
        raw_payload = part.get_payload()
        return raw_payload if isinstance(raw_payload, str) else ""
    charset = part.get_content_charset()
    if charset:
        try:
            return payload.decode(charset)
        except (LookupError, UnicodeDecodeError):
            pass
    return _decode_text(payload)


def _strip_html(value: str) -> str:
    collapsed = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", value)
    collapsed = re.sub(r"(?i)<br\s*/?>", "\n", collapsed)
    collapsed = re.sub(r"(?i)</p>", "\n", collapsed)
    collapsed = re.sub(r"(?s)<[^>]+>", " ", collapsed)
    collapsed = html.unescape(collapsed)
    collapsed = re.sub(r"[ \t]+", " ", collapsed)
    collapsed = re.sub(r"\n{3,}", "\n\n", collapsed)
    return collapsed.strip()


def _clean_header(value: str | None) -> str:
    if not value:
        return ""
    return " ".join(str(value).split())


def _format_addresses(values: list[str]) -> str:
    addresses = []
    for name, address in getaddresses(values):
        name = name.strip()
        address = address.strip()
        if name and address:
            addresses.append(f"{name} <{address}>")
        elif address:
            addresses.append(address)
        elif name:
            addresses.append(name)
    return ", ".join(addresses)


def _normalize_email_datetime(value: str | None) -> str | None:
    if not value:
        return None
    try:
        return parsedate_to_datetime(value).isoformat()
    except (TypeError, ValueError, IndexError):
        return _clean_header(value) or None


def extract_generic_financial_fields(text: str) -> dict[str, object]:
    invoice_numbers = list(dict.fromkeys(_INVOICE_NO_RE.findall(text or "")))[:10]
    dates = list(dict.fromkeys(_DATE_RE.findall(text or "")))[:20]
    amounts = list(dict.fromkeys(match.group(1).strip() for match in _AMOUNT_RE.finditer(text or "")))[:20]

    doc_type = None
    lowered = (text or "").lower()
    if "factuur" in lowered or "invoice" in lowered:
        doc_type = "invoice"
    elif "receipt" in lowered or "bon" in lowered:
        doc_type = "receipt"

    return {
        "document_type_hint": doc_type,
        "invoice_numbers": invoice_numbers,
        "dates": dates,
        "amounts": amounts,
    }
